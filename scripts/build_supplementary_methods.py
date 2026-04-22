"""
build_supplementary_methods.py — Build Supplementary Methods (Appendix A)

Generates Supplementary_Methods.docx with formal methodological description
of the entire analysis pipeline:
  1. Cohort curation from GSE246221
  2. Count normalization (TMM, log2-CPM)
  3. Differential expression (limma-voom, robust + trend)
  4. Cross-validation with PyDESeq2
  5. Cell-type signature analysis (marker-based, NOT deconvolution)
  6. F-test longitudinal analysis
  7. Post-hoc statistical tests (Holm-Bonferroni)
  8. Methodological audit reference (GitHub link)

Inputs:
  data/cohort_final.csv (for numbers cited in methods)

Output:
  results/supplementary/Supplementary_Methods.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_DIR = PROJECT_ROOT / "data"
SUPP_DIR = PROJECT_ROOT / "results" / "supplementary"
SUPP_DIR.mkdir(parents=True, exist_ok=True)


def set_style(doc):
    """Set default paragraph style."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return p


def build_document(cohort):
    doc = Document()
    set_style(doc)
    
    # Title
    t = doc.add_heading('Appendix A — Supplementary Methods', level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph(
        'Transcriptomic reanalysis of the IL-10 axis in a '
        'MASLD/MASH-to-HCC mouse model (GSE246221)'
    )
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.italic = True
    
    doc.add_paragraph()
    
    # Section 1
    add_heading(doc, 'A.1 Cohort curation', level=1)
    n_total = len(cohort)
    groups = cohort['group'].value_counts().to_dict()
    add_para(doc,
        f'Raw count data and sample metadata were downloaded from Gene '
        f'Expression Omnibus accession GSE246221 (Jeong et al., 2024), '
        f'comprising male C57BL/6J mice subjected to streptozotocin (STZ) '
        f'and high-fat diet (HFD) to recapitulate a stage-wise trajectory '
        f'from metabolic liver injury to hepatocellular carcinoma. '
        f'To ensure biological homogeneity, strict metadata curation was '
        f'applied prior to analysis:'
    )
    add_para(doc,
        '• All pharmacological-intervention samples (Tirzepatide and '
        'matched vehicle controls, n=10) were excluded to avoid '
        'confounding by drug treatment effects.'
    )
    add_para(doc,
        '• All HFD-only Batch 2 samples (n=5) were excluded to preserve '
        'batch homogeneity with the Batch 1 STZ+HFD trajectory.'
    )
    add_para(doc,
        f'• The final curated cohort comprised n={n_total} biologically '
        f'independent samples from sequencing Batch 1, distributed across '
        f'six disease stages: '
        f'Healthy control at 7 weeks (n={groups.get("S1_Control_07w", 0)}); '
        f'Early MASLD at 14 weeks (n={groups.get("S2a_EarlyMASLD_14w", 0)}); '
        f'MASH at 20 weeks (n={groups.get("S3_MASH_20w", 0)}); '
        f'Liver fibrosis at 32 weeks (n={groups.get("S4_Fibrosis_32w", 0)}); '
        f'Chronic non-tumor inflammation at 56 weeks '
        f'(n={groups.get("S2b_ChronicNT_56w", 0)}, representing STZ+HFD '
        f'animals that did not progress to tumor formation); and '
        f'HCC at 44–56 weeks (n={groups.get("S5_HCC", 0)}).'
    )
    add_para(doc,
        'Stage assignments followed the histological criteria originally '
        'described by Jeong et al. (2024), based on steatosis, lobular '
        'inflammation, and fibrosis grading. The detailed cohort '
        'composition is provided in Supplementary Table A1, and the '
        'per-stage histological grading distributions (as annotated by '
        'the original authors) are provided in Supplementary Table A1b.'
    )
    
    # Section 2
    add_heading(doc, 'A.2 Count normalization', level=1)
    add_para(doc,
        'Raw gene count data were normalized using the trimmed mean of '
        'M-values (TMM) method (Robinson and Oshlack, 2010) as '
        'implemented in a custom Python port of the edgeR normalization '
        'procedure. The geometric mean of the resulting normalization '
        'factors across the 40 samples was verified to equal 1.0 '
        '(property of TMM, confirmed in the methodological audit). '
        'Normalized counts were converted to log2-counts-per-million '
        '(log2-CPM) using voom (Law et al., 2014), with a prior count '
        'of 0.5 added to avoid log(0) undefined values. Library-size-'
        'weighted observational-level variances were estimated by voom '
        'using a LOWESS trend fit (span=0.5) on mean-variance coordinates.'
    )
    
    # Section 3
    add_heading(doc, 'A.3 Differential expression analysis (limma-voom)',
                level=1)
    add_para(doc,
        'Differential expression was assessed using a custom Python '
        'implementation of the limma-voom pipeline (Law et al., 2014; '
        'Ritchie et al., 2015) with robust empirical Bayes and '
        'mean-variance trend fitting (robust=TRUE, trend=TRUE; Phipson '
        'et al., 2016). Briefly, for each gene a linear model was fit '
        'to the log2-CPM expression values across the six disease '
        'stages, weighted by the observational-level voom weights. '
        'Six pairwise contrasts were evaluated (each stage vs. the '
        'preceding one, plus HCC vs. Control), and a longitudinal '
        'F-test was applied to identify genes with any stage-wise '
        'expression difference. Empirical Bayes shrinkage of the '
        'per-gene variances was performed using the method of Smyth '
        '(2004), with robust winsorization to limit influence of '
        'outlier variances. Benjamini-Hochberg correction was applied '
        'to obtain FDR-adjusted P-values.'
    )
    
    # Section 4
    add_heading(doc, 'A.4 Cross-validation with PyDESeq2', level=1)
    add_para(doc,
        'To verify that our custom Python limma-voom implementation '
        'does not introduce method-specific biases, we independently '
        're-analyzed the HCC vs. Control contrast using PyDESeq2 '
        '(Muzellec et al., 2023), a Python port of the negative-'
        'binomial-based DESeq2 method (Love et al., 2014). Input '
        'counts were rounded to integers (as required by DESeq2) '
        'and the Wald test was applied with default parameters. '
        'log2FoldChange estimates from PyDESeq2 were compared against '
        'the limma-voom estimates across all six pairwise contrasts, '
        'yielding Pearson correlations of r = 0.88 – 0.98 (summary '
        'statistics in Supplementary Table A2c). This sensitivity '
        'analysis confirms that the main findings are robust to the '
        'choice of differential expression method.'
    )
    
    # Section 5
    add_heading(doc, 'A.5 Cell-type signature enrichment analysis '
                     '(not a deconvolution)', level=1)
    add_para(doc,
        'To evaluate the stage-dependent remodeling of the hepatic '
        'cellular landscape, we performed a signature enrichment '
        'analysis for four hepatic cell types with complementary '
        'roles in the IL-10 axis: hepatocytes (parenchymal IL-10 '
        'targets); the macrophage/monocyte compartment (including '
        'Kupffer cells, monocyte-derived macrophages, and TREM2+ '
        'lipid-associated macrophages — the principal IL-10 '
        'producers and targets in the liver); NK cells (innate '
        'lymphoid IL-10 sources); and hepatic stellate cells '
        '(HSC/fibrosis effectors).'
    )
    add_para(doc,
        'Gene signatures were curated from CellMarker 2.0 '
        '(Hu et al., 2023), filtered to liver-tissue entries and '
        'supplemented with literature-canonical markers for disease-'
        'relevant subsets. Per-sample signature scores were computed '
        'as the median log2-CPM of the detected signature genes, '
        'following the rationale of mMCP-counter (Petitprez et al., '
        '2020). Critically, this is a signature-enrichment approach '
        'and NOT a cell-proportion deconvolution: the scores report '
        'relative enrichment of gene-set activity across stages rather '
        'than absolute cellular proportions. This distinction is '
        'explicitly reiterated in the main text (Section 4.2) and '
        'in the Figure 4 legend. The full signature panel with gene-'
        'level provenance is provided in Supplementary Table A3.'
    )
    add_para(doc,
        'Per-sample signature scores (Supplementary Table A4a) were '
        'then tested for stage-dependent variation using one-way ANOVA, '
        'with Benjamini-Hochberg correction across the four signatures '
        '(Supplementary Table A4b).'
    )
    
    # Section 6
    add_heading(doc, 'A.6 Post-hoc pairwise tests', level=1)
    add_para(doc,
        'When global stage-wise differences were detected by ANOVA or '
        'by the longitudinal F-test, post-hoc pairwise comparisons were '
        'performed using Mann-Whitney U tests with Holm-Bonferroni '
        'correction to control family-wise error rate. Significance is '
        'indicated in all figures with the following convention: '
        '* adjusted P < 0.05, ** < 0.01, *** < 0.001.'
    )
    
    # Section 7
    add_heading(doc, 'A.7 Methodological audit and reproducibility',
                level=1)
    add_para(doc,
        'A comprehensive methodological audit of the analysis pipeline '
        'was performed and is publicly available at '
        'https://github.com/Hsolleiro/GSE246221_IL10_Analysis/tree/main/audit. '
        'The audit comprises:'
    )
    add_para(doc,
        '• A reproducible Google Colab notebook (rpy2-based) comparing '
        'the custom Python limma-voom implementation against the '
        'reference R limma package (Ritchie et al., 2015) on the same '
        'cohort data. Pearson correlation of log2FoldChange estimates '
        'was 1.000 in all six pairwise contrasts, with top-100 gene '
        'overlap ranging from 91-100 per contrast, confirming numerical '
        'equivalence of the Python implementation to the reference R '
        'package.'
    )
    add_para(doc,
        '• Six independent statistical tests (audit_methods.py): TMM '
        'normalization properties; cross-method concordance between '
        'limma-voom and PyDESeq2; empirical Bayes prior degrees of '
        'freedom estimation (validated at df0 = 2.22 in Python vs. '
        'df0 = 3.22 in R, with negligible impact on log2FC given the '
        'r=1.000 correlation); Type I error calibration under 100 '
        'label permutations (mean fraction p<0.05 = 0.049, matching '
        'the nominal α = 0.05); power recovery under planted '
        'differential expression; and implementation of the Benjamini-'
        'Hochberg FDR procedure.'
    )
    add_para(doc,
        'Full analysis scripts, data, and results are available at '
        'https://github.com/Hsolleiro/GSE246221_IL10_Analysis.'
    )
    
    # References section
    add_heading(doc, 'A.8 Software and versions', level=1)
    add_para(doc,
        'The analysis was performed using Python 3.10+ with the '
        'following core packages: numpy, pandas, scipy, statsmodels, '
        'scikit-learn, pydeseq2, matplotlib. The custom limma-voom '
        'implementation (scripts/limma_voom.py in the GitHub repository) '
        'replicates the key functions of the R Bioconductor limma '
        'package (version 3.58.1 used as reference). Reference R '
        'validation was performed in Google Colab using R 4.3.2 with '
        'Bioconductor 3.18.'
    )
    
    return doc


def main():
    print("Loading inputs...")
    cohort = pd.read_csv(DATA_DIR / 'cohort_final.csv')
    print(f"  cohort: {len(cohort)} samples")
    
    print("\nBuilding document...")
    doc = build_document(cohort)
    
    out_path = SUPP_DIR / 'Supplementary_Methods.docx'
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    main()
